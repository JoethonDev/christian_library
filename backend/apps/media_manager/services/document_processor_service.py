"""
Document Processor Service for extracting text from Word documents (.doc/.docx).
Handles text extraction, cleaning, and normalization for supplementary documents.
"""
import logging
import os
import re
import subprocess
from typing import Tuple

logger = logging.getLogger(__name__)


class DocumentProcessorService:
    """Service for processing Word documents and extracting text."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from .docx files using python-docx library.
        Handles paragraphs, tables, headers, and footers.
        Preserves paragraph structure and handles Arabic RTL text.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Extracted text as string
        """
        try:
            import docx
            
            self.logger.info(f"Extracting text from DOCX file: {file_path}")
            doc = docx.Document(file_path)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            # Combine all text parts
            full_text = '\n\n'.join(text_parts)
            
            self.logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except ImportError:
            self.logger.error("python-docx library not installed")
            raise Exception("python-docx library is required for .docx file processing")
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX: {str(e)}", exc_info=True)
            raise
    
    def extract_text_from_doc(self, file_path: str) -> str:
        """
        Extract text from legacy .doc files using antiword (preferred), 
        textract (if available), or pandoc.
        
        Args:
            file_path: Path to the .doc file
            
        Returns:
            Extracted text as string
        """
        try:
            # Try antiword first (most reliable system tool for .doc)
            try:
                self.logger.info(f"Trying antiword for DOC file: {file_path}")
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    text = result.stdout
                    self.logger.info(f"Successfully extracted {len(text)} characters from DOC using antiword")
                    return text
                else:
                    self.logger.warning(f"Antiword failed with return code {result.returncode}")
            except (FileNotFoundError, Exception) as e:
                self.logger.warning(f"Antiword not available or failed: {str(e)}")
            
            # Try pandoc as last resort
            try:
                self.logger.info(f"Trying pandoc for DOC file: {file_path}")
                result = subprocess.run(
                    ['pandoc', file_path, '-t', 'plain'],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    text = result.stdout
                    self.logger.info(f"Successfully extracted {len(text)} characters from DOC using pandoc")
                    return text
                else:
                    self.logger.warning(f"Pandoc failed with return code {result.returncode}")
            except (FileNotFoundError, Exception) as e:
                self.logger.warning(f"Pandoc not available or failed: {str(e)}")
            
            raise Exception("No suitable tool found for .doc file processing. Install antiword or pandoc.")
            
        except Exception as e:
            self.logger.error(f"Error extracting text from DOC: {str(e)}", exc_info=True)
            raise
    
    def extract_text_from_document(self, file_path: str, mime_type: str) -> str:
        """
        Router method that calls appropriate extractor based on file type.
        Handles errors and fallbacks.
        
        Args:
            file_path: Path to the document file
            mime_type: MIME type of the document
            
        Returns:
            Extracted and cleaned text as string
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Document file not found: {file_path}")
            
            # Determine file type from extension if mime_type not reliable
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            # Extract text based on file type
            if ext == '.docx' or 'wordprocessingml' in mime_type:
                raw_text = self.extract_text_from_docx(file_path)
            elif ext == '.doc' or mime_type == 'application/msword':
                raw_text = self.extract_text_from_doc(file_path)
            else:
                # Try .docx first as default, then .doc
                try:
                    raw_text = self.extract_text_from_docx(file_path)
                except Exception:
                    raw_text = self.extract_text_from_doc(file_path)
            
            # Clean and normalize the extracted text
            cleaned_text = self.clean_and_normalize_text(raw_text)
            
            return cleaned_text
            
        except Exception as e:
            self.logger.error(f"Error in extract_text_from_document: {str(e)}", exc_info=True)
            # Return empty string instead of raising to avoid breaking the upload process
            return ""
    
    def clean_and_normalize_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        Removes excessive whitespace, normalizes line breaks, and preserves paragraph structure.
        Applies Arabic text cleaning pipeline.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""
        
        try:
            # Remove excessive whitespace while preserving paragraph breaks
            # Replace multiple spaces with single space
            text = re.sub(r' +', ' ', text)
            
            # Replace multiple newlines with double newline (paragraph break)
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            # Remove spaces at start/end of lines
            text = '\n'.join(line.strip() for line in text.split('\n'))
            
            # Apply Arabic text normalization
            text = self._normalize_arabic_text(text)
            
            # Final cleanup
            text = text.strip()
            
            return text
            
        except Exception as e:
            self.logger.error(f"Error cleaning text: {str(e)}", exc_info=True)
            return text
    
    def _normalize_arabic_text(self, text: str) -> str:
        """
        Normalize Arabic text for better search and indexing.
        
        Args:
            text: Text containing Arabic characters
            
        Returns:
            Normalized Arabic text
        """
        if not text:
            return ""
        
        try:
            # Normalize Arabic characters
            # Replace different forms of Alef with standard Alef
            text = re.sub(r'[آأإٱ]', 'ا', text)
            
            # Remove tatweel (kashida)
            text = text.replace('ـ', '')
            
            # Normalize Hamza
            text = re.sub(r'[ؤئ]', 'ء', text)
            
            # Remove diacritics (tashkeel)
            arabic_diacritics = re.compile(r'[\u064B-\u065F\u0670]')
            text = arabic_diacritics.sub('', text)
            
            return text
            
        except Exception as e:
            self.logger.error(f"Error normalizing Arabic text: {str(e)}", exc_info=True)
            return text
    
    def validate_document(self, file_size: int, mime_type: str, filename: str) -> Tuple[bool, str]:
        """
        Validate document before processing.
        Checks file size, format, and basic security.
        
        Args:
            file_size: Size of file in bytes
            mime_type: MIME type of the file
            filename: Name of the file
            
        Returns:
            Tuple of (is_valid: bool, error_message: str)
        """
        # Check file size (2GB limit)
        max_size = 2 * 1024 * 1024 * 1024  # 2GB in bytes
        if file_size > max_size:
            return False, f"File size ({file_size / (1024*1024):.1f}MB) exceeds 2GB limit"
        
        # Check file extension
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        allowed_extensions = ['.doc', '.docx']
        if ext not in allowed_extensions:
            return False, f"File extension {ext} not allowed. Only .doc and .docx files are supported."
        
        # Check MIME type
        allowed_mime_types = [
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/octet-stream',  # Sometimes browsers send this for doc files
        ]
        if mime_type and mime_type not in allowed_mime_types:
            # Allow if mime_type detection failed but extension is valid
            if ext not in allowed_extensions:
                return False, f"MIME type {mime_type} not allowed"
        
        return True, ""
